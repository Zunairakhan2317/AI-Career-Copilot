"""
ATS Tailoring Service

Takes a stored resume + a job description, asks Gemini to rewrite
experience bullets for ATS compliance, and returns both structured
JSON (for the frontend preview) and a .docx file (for download).
"""

from __future__ import annotations
import io
import json
import re
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database import supabase
from llm_client import call_llm


def get_resume_for_user(resume_id: str, user_id: str) -> dict | None:
    """Fetch a single resume, scoped to the user (security)."""
    response = (
        supabase.table("resumes")
        .select("*")
        .eq("id", resume_id)
        .eq("user_id", user_id)
        .execute()
    )
    return response.data[0] if response.data else None


def generate_tailored_resume(resume_id: str, user_id: str, job_description: str) -> dict:
    """
    Rewrite the resume's experience bullets to align with the JD,
    and return structured content + a base64-encoded .docx file.
    """
    if not job_description.strip():
        raise ValueError("Job description cannot be empty.")

    resume = get_resume_for_user(resume_id, user_id)
    if not resume:
        raise ValueError("Resume not found for this user.")

    parsed = resume.get("parsed_data") or {}
    if not parsed:
        raise ValueError("Resume exists but contains no parsed data.")

    # 1. Build the LLM prompt
    prompt = f"""You are an expert resume writer specializing in ATS-optimized resumes.

Given the candidate's parsed resume and a target job description, rewrite the
candidate's experience bullet points to maximize ATS match for the role.

Resume (JSON):
{json.dumps(parsed, indent=2)}

Target Job Description:
{job_description}

Return ONLY a valid JSON object with this exact structure (no markdown, no commentary):
{{
  "rewritten_summary": "A tailored 2-3 sentence professional summary rewritten for this role.",
  "experience": [
    {{
      "original_title": "...",
      "company": "...",
      "start_date": "...",
      "end_date": "...",
      "is_current": true|false,
      "rewritten_bullets": [
        "Action verb + task + measurable result using keywords from the JD",
        "..."
      ]
    }}
  ],
  "skills_to_emphasize": ["skill1", "skill2"],
  "keywords_added": ["keyword1", "keyword2"],
  "ats_match_estimate": 0-100
}}

RULES:
- Keep the same number of experience entries as the original.
- Each experience entry should have 3-5 rewritten bullets.
- Pull keywords and required skills directly from the job description.
- Use strong action verbs and quantify results wherever possible.
- Do NOT invent experience, companies, or dates that aren't in the original.
- ats_match_estimate must be an integer from 0 to 100.
"""

    # 2. Call the LLM
    try:
        llm_response = call_llm(prompt, provider="auto")
    except Exception as e:
        raise RuntimeError(f"LLM tailoring failed: {e}")

    # 3. Parse the response
    cleaned = re.sub(r"^```(?:json)?\s*", "", llm_response.strip())
    cleaned = re.sub(r"\s*```$", "", cleaned).strip()
    try:
        tailored = json.loads(cleaned)
    except json.JSONDecodeError:
        raise RuntimeError("LLM returned invalid JSON. Please try again.")

    # 4. Validate minimum required fields
    if "rewritten_summary" not in tailored or "experience" not in tailored:
        raise RuntimeError("LLM response missing required fields.")

    # 5. Generate the .docx file
    docx_bytes = build_docx(parsed, tailored)

    import base64
    return {
        "resume_id": resume_id,
        "target_role": tailored.get("rewritten_summary", "")[:80],
        "tailored": tailored,
        "docx_base64": base64.b64encode(docx_bytes).decode("utf-8"),
        "docx_filename": f"tailored_resume_{resume_id[:8]}.docx",
    }


def build_docx(parsed: dict, tailored: dict) -> bytes:
    """
    Build a .docx file from the tailored content.
    Layout: Contact → Summary → Experience (with rewritten bullets) → Skills.
    """
    doc = Document()

    # Set base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    contact = parsed.get("contact") or {}

    # --- Header: Name + Contact ---
    name = contact.get("full_name") or "Your Name"
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(20)

    contact_parts = [
        contact.get("email"),
        contact.get("phone"),
        contact.get("location"),
    ]
    contact_parts = [p for p in contact_parts if p]
    if contact_parts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(" • ".join(contact_parts))
        cr.font.size = Pt(10)

    # Links row
    link_parts = []
    if contact.get("linkedin"):
        link_parts.append(("LinkedIn", contact["linkedin"]))
    if contact.get("github"):
        link_parts.append(("GitHub", contact["github"]))
    if contact.get("portfolio"):
        link_parts.append(("Portfolio", contact["portfolio"]))
    if link_parts:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, (label, url) in enumerate(link_parts):
            if i > 0:
                lp.add_run(" • ").font.size = Pt(10)
            r = lp.add_run(label)
            r.font.size = Pt(10)
            r.font.color.rgb = None  # default

    # --- Professional Summary ---
    summary_text = tailored.get("rewritten_summary") or parsed.get("summary")
    if summary_text:
        doc.add_paragraph()  # spacer
        h = doc.add_paragraph()
        hr = h.add_run("PROFESSIONAL SUMMARY")
        hr.bold = True
        hr.font.size = Pt(12)
        sp = doc.add_paragraph(summary_text)
        sp.paragraph_format.space_after = Pt(6)

    # --- Experience ---
    tailored_exp = tailored.get("experience") or []
    if tailored_exp:
        doc.add_paragraph()
        h = doc.add_paragraph()
        hr = h.add_run("PROFESSIONAL EXPERIENCE")
        hr.bold = True
        hr.font.size = Pt(12)

        for entry in tailored_exp:
            # Title + company line
            title = entry.get("original_title") or entry.get("title") or "Position"
            company = entry.get("company") or "Company"
            start = entry.get("start_date") or ""
            end = entry.get("end_date") or ("Present" if entry.get("is_current") else "")

            tp = doc.add_paragraph()
            tr = tp.add_run(f"{title} — {company}")
            tr.bold = True
            tr.font.size = Pt(11)

            if start or end:
                dp = doc.add_paragraph()
                dr = dp.add_run(f"{start} – {end}".strip(" –"))
                dr.italic = True
                dr.font.size = Pt(10)
                dr.font.color.rgb = None

            # Bullets
            for bullet in entry.get("rewritten_bullets", []):
                bp = doc.add_paragraph(bullet, style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)

    # --- Skills to emphasize ---
    skills = tailored.get("skills_to_emphasize") or []
    if not skills:
        # Fall back to original skills if LLM didn't return any
        skills = [
            s.get("name")
            for s in (parsed.get("skills") or [])
            if s.get("name")
        ]
    if skills:
        doc.add_paragraph()
        h = doc.add_paragraph()
        hr = h.add_run("SKILLS")
        hr.bold = True
        hr.font.size = Pt(12)
        doc.add_paragraph(", ".join(skills))

    # --- Education (from original parsed data) ---
    education = parsed.get("education") or []
    if education:
        doc.add_paragraph()
        h = doc.add_paragraph()
        hr = h.add_run("EDUCATION")
        hr.bold = True
        hr.font.size = Pt(12)
        for edu in education:
            institution = edu.get("institution") or "Institution"
            degree = edu.get("degree") or ""
            field = edu.get("field_of_study") or ""
            dates = " – ".join(
                [d for d in [edu.get("start_date"), edu.get("end_date")] if d]
            )
            ep = doc.add_paragraph()
            er = ep.add_run(institution)
            er.bold = True
            er.font.size = Pt(11)
            if degree or field:
                ep.add_run(f"  |  {degree}{' in ' + field if field else ''}").font.size = Pt(10)
            if dates:
                dp = doc.add_paragraph(dates)
                dp.runs[0].italic = True
                dp.runs[0].font.size = Pt(10)

    # --- Certifications (from original) ---
    certs = parsed.get("certifications") or []
    if certs:
        doc.add_paragraph()
        h = doc.add_paragraph()
        hr = h.add_run("CERTIFICATIONS")
        hr.bold = True
        hr.font.size = Pt(12)
        for c in certs:
            doc.add_paragraph(c, style="List Bullet")

    # --- Save to bytes ---
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
